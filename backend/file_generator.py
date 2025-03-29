from docx import Document
import pandas as pd

def generate_word_file(requirements, filename="requirements.docx"):
    doc = Document()
    doc.add_heading("Extracted Software Requirements", level=1)

    for category, reqs in requirements.items():
        doc.add_heading(category, level=2)
        for req in reqs:
            doc.add_paragraph(f"- {req}")

    doc.save(filename)
    return filename

def generate_excel_file(requirements, filename="user_stories.xlsx"):
    data = []
    for category, reqs in requirements.items():
        for req in reqs:
            priority = "Must Have" if "should" in req.lower() else "Should Have"
            data.append({"Category": category, "Requirement": req, "Priority": priority})

    df = pd.DataFrame(data)
    df.to_excel(filename, index=False)
    return filename
